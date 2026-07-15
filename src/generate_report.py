# -*- coding: utf-8 -*-
"""
生成中华麋鹿园评价数据分析 Word 报告
==========================================================================
功能概述:
    使用 python-docx 库自动生成结构化的 Word 分析报告（.docx），
    报告包含封面、目录、8 个章节正文、11 张图表和 2 张汇总表格。

报告结构:
    封面        景区名称、数据来源、样本量、报告日期
    目录        八章目录列表
    第一章      执行摘要（核心结论 + 关键指标）
    第二章      数据概况（基本情况表 + 评分分布表）
    第三章      评分分析（总体分布 + 分项维度对比 + 图 01-02）
    第四章      时间趋势分析（年度 + 月度 + 季节性 + 图 03-05）
    第五章      用户画像分析（等级分布 + 地区分布 + 图 06-07）
    第六章      文本情感分析（词云 + 高频词 + 好评差评对比 + 图 08-11）
    第七章      主要发现与改进建议
    第八章      附录（数据采集说明 + 分析方法）

数据源:    config.REVIEWS_CSV + config.CHART_DIR 下的 11 张 PNG 图表
输出路径:  config.DOCX_REPORT

"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

import config
import data_utils

# 常用颜色（RGB 元组），供正文样式统一引用
BRAND_BLUE = (46, 86, 139)   # 品牌蓝：标题、强调
DARK_GRAY = (46, 46, 46)     # 深灰：一级/二级标题
MID_GRAY = (80, 80, 80)      # 中灰：三级标题
BODY_GRAY = (60, 60, 60)     # 正文灰


# ========== 辅助函数：中文字体与样式设置 ==========

def set_chinese_font(run, font_name="Microsoft YaHei", size=12, bold=False, color=None):
    """
    设置 Word 文本 Run 的中文字体、字号、粗体和颜色。

    python-docx 默认不设置东亚字体（w:eastAsia），导致中文显示为宋体。
    本函数通过 XML 操作显式设置东亚字体属性，确保中文字体正确应用。

    Args:
        run: docx Run 对象
        font_name (str): 字体名称，默认 "Microsoft YaHei"（微软雅黑）
        size (int): 字号（磅），默认 12
        bold (bool): 是否粗体
        color (tuple|None): RGB 颜色元组 (R, G, B)，None 表示默认黑色
    """
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    # 设置东亚字体属性（关键步骤：确保中文字符使用指定字体）
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        font.color.rgb = RGBColor(*color)


def add_heading_cn(doc, text, level=1):
    """添加带中文字体的标题段落，字号与颜色随级别自动调整。"""
    # 各级别对应的字号梯度
    size_map = {1: 20, 2: 16, 3: 14}
    size = size_map.get(level, 12)
    color = DARK_GRAY if level <= 2 else MID_GRAY

    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_chinese_font(run, size=size, bold=True, color=color)
    return p


def add_paragraph_cn(doc, text, size=12, bold=False, color=BODY_GRAY,
                     alignment=None, first_line_indent=False):
    """
    添加带中文样式的正文段落（1.5 倍行距、段后 8pt）。

    Args:
        doc: Document 对象
        text (str): 段落文本
        size (int): 字号（磅）
        bold (bool): 是否粗体
        color (tuple): RGB 颜色元组
        alignment: 对齐方式（WD_ALIGN_PARAGRAPH 枚举），None 为默认左对齐
        first_line_indent (bool): 是否首行缩进 0.3 英寸
    """
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet_cn(doc, text, size=11):
    """添加带中文字体的项目符号段落（Word 内置 List Bullet 样式）。"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_chinese_font(run, size=size, color=BODY_GRAY)
    return p


def add_chart(doc, chart_name, width=6.0):
    """
    在文档中插入图表图片（居中对齐）。

    Args:
        doc: Document 对象
        chart_name (str): 图表文件名（不含扩展名），如 "01_评分分布"
        width (float): 图片宽度（英寸），默认 6.0

    Returns:
        Paragraph|None: 包含图片的段落，若文件不存在则返回 None
    """
    path = config.CHART_DIR / f"{chart_name}.png"
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        return p
    return None


def add_table_summary(doc, headers, rows):
    """
    添加带样式的汇总表格（Light Grid Accent 1，表头白字加粗）。

    Args:
        doc: Document 对象
        headers (list): 表头列表
        rows (list): 数据行列表，每行是一个列表
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # 表头行
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_chinese_font(run, size=11, bold=True, color=(255, 255, 255))

    # 数据行
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    set_chinese_font(run, size=11, color=BODY_GRAY)
    return table


# ========== 关键指标计算 ==========

def compute_metrics(df, df_valid):
    """
    计算报告正文与表格引用的所有关键指标（全部由数据计算，避免硬编码漂移）。

    Returns:
        dict: 指标名 -> 数值
    """
    total = len(df)
    score_5 = int((df["总评分"] == 5.0).sum())
    score_4 = int((df["总评分"] == 4.0).sum())
    score_3 = int((df["总评分"] == 3.0).sum())
    score_2 = int((df["总评分"] == 2.0).sum())
    score_1 = int((df["总评分"] == 1.0).sum())
    negative = int((df["总评分"] <= 2.0).sum())  # 差评（1-2星）

    # 分项维度均分
    dim_scenery = df["景色评分"].mean()
    dim_fun = df["趣味评分"].mean()
    dim_value = df["性价比评分"].mean()

    # 季节性（按春夏秋冬顺序）
    seas = (df_valid.groupby("季节").agg(n=("评论ID", "count"))
            .reindex(config.SEASON_ORDER).reset_index())
    seasonal = {row["季节"]: (int(row["n"]), round(row["n"] / total * 100, 1))
                for _, row in seas.iterrows()}

    # 用户等级 TOP3
    mem = df["用户等级"].value_counts()
    member = {k: (int(mem.get(k, 0)), round(mem.get(k, 0) / total * 100, 1))
              for k in ["铂金贵宾", "黄金贵宾", "钻石贵宾"]}
    vip_rate = round(sum(v[0] for v in member.values()) / total * 100, 1)

    # IP 属地 TOP4 + 长三角占比
    ip = df["IP属地"].value_counts()
    ip_top = {k: (int(ip.get(k, 0)), round(ip.get(k, 0) / total * 100, 1))
              for k in ["江苏", "上海", "浙江", "山东"]}
    yrd_rate = round((ip.get("江苏", 0) + ip.get("上海", 0) + ip.get("浙江", 0))
                     / total * 100, 1)

    # 年度峰值
    yr = df_valid.groupby("年份").agg(n=("评论ID", "count")).reset_index()
    peak_year = int(yr.sort_values("n", ascending=False).iloc[0]["年份"])
    year_2025 = int(yr[yr["年份"] == 2025]["n"].iloc[0]) if 2025 in yr["年份"].values else 0
    year_2026 = int(yr[yr["年份"] == 2026]["n"].iloc[0]) if 2026 in yr["年份"].values else 0

    # 情感极性（SnowNLP 模型）
    data_utils.add_sentiment_polarity(df)
    model_pos = int((df["情感分类"] == "正面").sum())
    model_neu = int((df["情感分类"] == "中性").sum())
    model_neg = int((df["情感分类"] == "负面").sum())
    model_pos_rate = round(model_pos / total * 100, 1)
    model_neu_rate = round(model_neu / total * 100, 1)
    model_neg_rate = round(model_neg / total * 100, 1)
    agree_rate = round(
        ((df["总评分"] >= config.POSITIVE_MIN_SCORE) & (df["情感分类"] == "正面")).sum()
        / total * 100, 1)
    silent_neg = int(
        ((df["总评分"] >= config.POSITIVE_MIN_SCORE) & (df["情感分类"] == "负面")).sum())
    silent_pos = int(
        ((df["总评分"] <= config.NEGATIVE_MAX_SCORE) & (df["情感分类"] == "正面")).sum())
    model_neg_words = [w for w, _ in
                       data_utils.word_frequency(
                           df[df["情感分类"] == "负面"]["评论内容"].dropna()).most_common(10)]
    model_pos_words = [w for w, _ in
                       data_utils.word_frequency(
                           df[df["情感分类"] == "正面"]["评论内容"].dropna()).most_common(10)]

    return {
        "total": total,
        "avg_score": df["总评分"].mean(),
        "score_5": score_5,
        "score_4": score_4,
        "score_3": score_3,
        "score_2": score_2,
        "score_1": score_1,
        "positive": score_5 + score_4,
        "negative": negative,
        "positive_rate": (score_5 + score_4) / total * 100,
        "negative_rate": negative / total * 100,
        "image_rate": (df["图片数"] > 0).sum() / total * 100,
        "dim_scenery": dim_scenery,
        "dim_fun": dim_fun,
        "dim_value": dim_value,
        "seasonal": seasonal,
        "member": member,
        "vip_rate": vip_rate,
        "ip_top": ip_top,
        "yrd_rate": yrd_rate,
        "peak_year": peak_year,
        "year_2025": year_2025,
        "year_2026": year_2026,
        "model_pos": model_pos,
        "model_neu": model_neu,
        "model_neg": model_neg,
        "model_pos_rate": model_pos_rate,
        "model_neu_rate": model_neu_rate,
        "model_neg_rate": model_neg_rate,
        "agree_rate": agree_rate,
        "silent_neg": silent_neg,
        "silent_pos": silent_pos,
        "model_neg_words": model_neg_words,
        "model_pos_words": model_pos_words,
        # 使用 df_valid 的评论字数（data_utils 已补全该列）
        "avg_words": df_valid["评论字数"].mean(),
        "date_min": df_valid["发布时间_dt"].min(),
        "date_max": df_valid["发布时间_dt"].max(),
        "year_min": df_valid["年份"].min(),
        "year_max": df_valid["年份"].max(),
    }


# ========== 各部分内容构建 ==========

def build_cover(doc, m):
    """封面：居中排版景区名称 + 标题 + 数据概况 + 报告日期。"""
    print("生成封面...")

    def centered(text, size, color, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_chinese_font(run, size=size, bold=bold, color=color)
        return p

    centered("\n\n\n", 12, (120, 120, 120))
    centered("中华麋鹿园", 32, BRAND_BLUE, bold=True)
    centered("携程评价数据分析报告", 28, BRAND_BLUE, bold=True)
    centered("\n\n\n", 12, (120, 120, 120))
    centered(
        f"景区: {config.SIGHT_FULL_NAME}\n数据来源: 携程旅行网\n"
        f"样本量: {m['total']} 条评价\n时间跨度: 2016-2026",
        14, MID_GRAY,
    )
    centered(f"\n\n\n报告日期: {datetime.now().strftime('%Y年%m月%d日')}", 12, (120, 120, 120))
    doc.add_page_break()


def build_toc(doc):
    """目录页（手动编排）。"""
    add_heading_cn(doc, "报告目录", level=1)
    for item in ["一、执行摘要", "二、数据概况", "三、评分分析", "四、时间趋势分析",
                 "五、用户画像分析", "六、文本情感分析", "七、主要发现与改进建议", "八、附录"]:
        add_bullet_cn(doc, item)
    doc.add_page_break()


def build_chapter_summary(doc, m):
    """第一章：执行摘要。"""
    add_heading_cn(doc, "一、执行摘要", level=1)
    add_paragraph_cn(doc, f"本报告基于携程平台上盐城中华麋鹿园的 {m['total']} 条游客评价数据，分析时间跨度为 2016 年至 2026 年。通过评分分布、时间趋势、用户画像、文本情感等多维度分析，为景区运营管理提供数据洞察和改进建议。")

    add_paragraph_cn(doc, "核心结论:", size=14, bold=True, color=BRAND_BLUE)
    add_bullet_cn(doc, f"景区整体口碑良好：平均评分 {m['avg_score']:.2f} 分，好评率（4-5星）达到 {m['positive_rate']:.1f}%。")
    add_bullet_cn(doc, "游客最认可的体验是：近距离喂麋鹿、乘坐观光车游览、亲子互动。")
    add_bullet_cn(doc, "主要短板是：排队等候时间长、性价比感知偏低、服务细节有待优化。")
    add_bullet_cn(doc, "客流旺季集中在夏季（4-6月），本地及周边省份游客是主力客群。")
    add_bullet_cn(doc, "2025年起评价量爆发式增长，景区热度显著提升，需加强高峰期管理。")
    doc.add_page_break()


def build_chapter_overview(doc, df, m):
    """第二章：数据概况（基本情况表 + 评分分布表）。"""
    add_heading_cn(doc, "二、数据概况", level=1)
    add_paragraph_cn(doc, f"本次分析共采集 {m['total']} 条有效评价数据，覆盖时间范围为 {m['date_min'].strftime('%Y-%m-%d')} 至 {m['date_max'].strftime('%Y-%m-%d')}。数据字段包括用户昵称、评分、分项评分（景色、趣味、性价比）、评论内容、发布时间、图片数量、IP属地等。")

    add_paragraph_cn(doc, "表1：数据基本情况", size=13, bold=True, color=MID_GRAY)
    add_table_summary(doc, ["指标", "数值"], [
        ["总评价数", f"{m['total']} 条"],
        ["平均评分", f"{m['avg_score']:.2f} / 5.0"],
        ["好评数（4-5星）", f"{m['positive']} 条 ({m['positive_rate']:.1f}%)"],
        ["中评数（3星）", f"{m['score_3']} 条 ({m['score_3']/m['total']*100:.1f}%)"],
        ["差评数（1-2星）", f"{m['negative']} 条 ({m['negative_rate']:.1f}%)"],
        ["有图评论", f"{(df['图片数']>0).sum()} 条 ({m['image_rate']:.1f}%)"],
        ["平均评论字数", f"{m['avg_words']:.0f} 字"],
        ["覆盖年份", f"{m['year_min']} - {m['year_max']}"],
    ])

    add_paragraph_cn(doc, "表2：评分分布明细", size=13, bold=True, color=MID_GRAY)
    t = m["total"]
    add_table_summary(doc, ["评分", "数量", "占比", "累计占比"], [
        ["5星", m["score_5"], f"{m['score_5']/t*100:.1f}%", f"{m['score_5']/t*100:.1f}%"],
        ["4星", m["score_4"], f"{m['score_4']/t*100:.1f}%", f"{(m['score_5']+m['score_4'])/t*100:.1f}%"],
        ["3星", m["score_3"], f"{m['score_3']/t*100:.1f}%", f"{(m['score_5']+m['score_4']+m['score_3'])/t*100:.1f}%"],
        ["2星", m["score_2"], f"{m['score_2']/t*100:.1f}%", "-"],
        ["1星", m["score_1"], f"{m['score_1']/t*100:.1f}%", "-"],
    ])
    doc.add_page_break()


def build_chapter_scores(doc, m):
    """第三章：评分分析（总体分布 + 分项维度）。"""
    add_heading_cn(doc, "三、评分分析", level=1)

    add_heading_cn(doc, "3.1 总体评分分布", level=2)
    add_paragraph_cn(doc, f"中华麋鹿园总体口碑较好，平均评分 {m['avg_score']:.2f} 分。5星评价最多，达到 {m['score_5']} 条（占比 {m['score_5']/m['total']*100:.1f}%），4星评价 {m['score_4']} 条（{m['score_4']/m['total']*100:.1f}%）。好评合计占比 {m['positive_rate']:.1f}%，差评占比仅 {m['negative_rate']:.1f}%。")
    add_chart(doc, "01_评分分布", width=5.5)

    add_heading_cn(doc, "3.2 分项评分对比", level=2)
    add_paragraph_cn(doc, f'从评分维度来看，游客对景区"景色"评分最高（{m["dim_scenery"]:.2f}分），其次是"趣味"（{m["dim_fun"]:.2f}分），"性价比"评分相对较低（{m["dim_value"]:.2f}分）。这表明景区自然景观和互动体验具有较强吸引力，但在价格感知、配套服务方面仍有提升空间。')
    add_chart(doc, "02_评分维度对比", width=5.5)

    add_paragraph_cn(doc, "主要洞察:", size=12, bold=True, color=BRAND_BLUE)
    add_bullet_cn(doc, f"景色（{m['dim_scenery']:.2f}分）：湿地生态、麋鹿成群是景区最大亮点，游客普遍认可。")
    add_bullet_cn(doc, f"趣味（{m['dim_fun']:.2f}分）：喂麋鹿、观光车等互动体验深受亲子家庭喜爱。")
    add_bullet_cn(doc, f"性价比（{m['dim_value']:.2f}分）：相对最薄弱，部分游客认为门票、交通项目价格偏高。")
    doc.add_page_break()


def build_chapter_time(doc, m):
    """第四章：时间趋势分析（年度 + 月度 + 季节性）。"""
    add_heading_cn(doc, "四、时间趋势分析", level=1)

    add_heading_cn(doc, "4.1 年度评价量趋势", level=2)
    add_paragraph_cn(doc, f"从年度维度看，2025年中华麋鹿园评价量出现爆发式增长，达到 {m['year_2025']} 条，远超往年。2026年截至当前已累积 {m['year_2026']} 条评价，热度持续上升。这与景区近年推广力度加大、亲子游市场增长等因素密切相关。")
    add_chart(doc, "03_年度趋势", width=6.0)

    add_heading_cn(doc, "4.2 月度评价趋势", level=2)
    add_paragraph_cn(doc, "从2023年至今的月度数据看，评价量呈现明显波动。2025年4-6月、2026年4-6月为评价高峰，与春季、初夏旅游旺季吻合。景区可在这些时段提前部署客流疏导、服务保障等资源。")
    add_chart(doc, "04_月度趋势", width=6.0)

    add_heading_cn(doc, "4.3 季节性特征", level=2)
    add_paragraph_cn(doc, f"夏季（4-6月）是评价最多的季节，共 {m['seasonal']['夏季(4-6月)'][0]} 条，占全年 {m['seasonal']['夏季(4-6月)'][1]}%。冬季（10-12月）次之，{m['seasonal']['冬季(10-12月)'][0]} 条；春季（1-3月）{m['seasonal']['春季(1-3月)'][0]} 条；秋季（7-9月）最少，{m['seasonal']['秋季(7-9月)'][0]} 条。各季节评分均在 4.4 分以上，整体波动不大。")
    add_chart(doc, "05_季节性分析", width=5.5)

    add_paragraph_cn(doc, "运营建议:", size=12, bold=True, color=BRAND_BLUE)
    add_bullet_cn(doc, "4-6月为绝对旺季，应提前增派观光车、讲解员、检票人员。")
    add_bullet_cn(doc, "7-9月（秋季）评价量最少，可结合暑期推出家庭套票、夜场等活动拉动客流。")
    add_bullet_cn(doc, "冬季（10-12月）有稳定客流，可围绕麋鹿繁殖季、湿地观鸟等主题开展营销。")
    doc.add_page_break()


def build_chapter_users(doc, m):
    """第五章：用户画像分析（等级分布 + 地区分布）。"""
    add_heading_cn(doc, "五、用户画像分析", level=1)

    add_heading_cn(doc, "5.1 用户等级分布", level=2)
    add_paragraph_cn(doc, f"从携程用户等级看，景区吸引了大量高等级用户：铂金贵宾占比最高（{m['member']['铂金贵宾'][1]}%），其次为黄金贵宾（{m['member']['黄金贵宾'][1]}%）和钻石贵宾（{m['member']['钻石贵宾'][1]}%）。高等级用户占比高，说明景区对资深旅行者、家庭出游群体具有较强吸引力。")
    add_chart(doc, "06_用户等级分布", width=5.5)

    add_heading_cn(doc, "5.2 来源地区分布", level=2)
    add_paragraph_cn(doc, f"江苏省游客是绝对主力，占比超过一半（{m['ip_top']['江苏'][1]}%），其次是上海（{m['ip_top']['上海'][1]}%）、浙江（{m['ip_top']['浙江'][1]}%）、山东（{m['ip_top']['山东'][1]}%）。景区属于典型的近程周边游目的地，长三角客群占主导地位。")
    add_chart(doc, "07_来源地区TOP10", width=5.5)

    add_paragraph_cn(doc, "营销建议:", size=12, bold=True, color=BRAND_BLUE)
    add_bullet_cn(doc, "继续深耕长三角市场，特别是上海、浙江的高铁/自驾游客群。")
    add_bullet_cn(doc, "针对高等级会员推出专属权益或家庭套餐，提升复购率。")
    add_bullet_cn(doc, "在江苏本地媒体、亲子类KOL渠道加强投放。")
    doc.add_page_break()


def build_chapter_text(doc, m):
    """第六章：文本情感分析（词云 + 高频词 + 好评差评 + 长度）。"""
    add_heading_cn(doc, "六、文本情感分析", level=1)

    add_heading_cn(doc, "6.1 整体词云", level=2)
    add_paragraph_cn(doc, f'通过 jieba 对 {m["total"]} 条评论内容进行分词统计，"观光车"、"胡萝卜"、"近距离"、"体验"、"孩子"等词汇出现频率最高。游客最关注的核心体验是：乘坐观光车、购买胡萝卜喂麋鹿、与孩子近距离互动。')
    add_chart(doc, "09_词云", width=5.5)

    add_heading_cn(doc, "6.2 高频关键词 TOP15", level=2)
    add_chart(doc, "10_高频词TOP15", width=5.5)

    add_heading_cn(doc, "6.3 好评 vs 差评关键词对比", level=2)
    add_paragraph_cn(doc, "将4-5星评价定义为好评，1-2星评价定义为差评，分别提取高频关键词。好评关键词集中在互动体验（观光车、胡萝卜、近距离、孩子），差评关键词集中在排队等候（排队、小时）、价格与服务（门票、工作人员、电瓶车）。")
    add_chart(doc, "11_好评差评关键词对比", width=6.0)

    add_heading_cn(doc, "6.4 评论长度分布", level=2)
    add_paragraph_cn(doc, f"游客平均评论字数为 {m['avg_words']:.0f} 字。大部分评论集中在 11-200 字之间，说明游客愿意分享具体体验，但长篇游记类评价相对较少。超过 1000 字的深度评价数量不多，但通常包含详细攻略和情感表达。")
    add_chart(doc, "08_评论长度分布", width=5.5)

    add_heading_cn(doc, "6.5 情感极性模型分析（SnowNLP）", level=2)
    add_paragraph_cn(doc, f"为弥补「仅按星级分段」的局限，本报告引入 SnowNLP 中文情感极性模型，对每条评论做 0-1 语义打分（正面≥0.6、负面≤0.4、中间为中性）。结果显示：正面 {m['model_pos_rate']}%、中性 {m['model_neu_rate']}%、负面 {m['model_neg_rate']}%。值得注意的是，模型判定的负面比例（{m['model_neg_rate']}%）显著高于按星级的差评率（{m['negative_rate']:.1f}%），说明不少高分评价在文字中仍夹杂负面体验（如排队、价格）。其中 {m['silent_neg']} 条 4-5 星评价被模型判为负面（隐性不满），值得在好评中进一步挖掘改进线索。")
    add_chart(doc, "12_情感极性分布", width=5.5)
    add_paragraph_cn(doc, "模型判定为负面评论的高频词（痛点信号）：" + "、".join(m["model_neg_words"][:10]) + "。")
    doc.add_page_break()


def build_chapter_findings(doc, m):
    """第七章：主要发现与改进建议。"""
    add_heading_cn(doc, "七、主要发现与改进建议", level=1)

    add_heading_cn(doc, "7.1 主要发现", level=2)
    add_bullet_cn(doc, f"口碑优势显著：平均 {m['avg_score']:.2f} 分、好评率 {m['positive_rate']:.1f}%，5A 级景区品牌形象得到游客认可。")
    add_bullet_cn(doc, "核心吸引力明确：近距离喂麋鹿、亲子互动、湿地观光是游客最满意的三大体验。")
    add_bullet_cn(doc, f'性价比是短板：分项评分中"性价比"最低（{m["dim_value"]:.2f}分），部分游客认为门票和项目价格偏高。')
    add_bullet_cn(doc, '排队是最大痛点：差评关键词中"排队"、"小时"出现频繁，高峰期游客等待体验不佳。')
    add_bullet_cn(doc, f"客流高度集中：夏季（4-6月）占全年 {m['seasonal']['夏季(4-6月)'][1]}%，景区淡旺季差异明显。")
    add_bullet_cn(doc, f"客群地域集中：江苏、上海、浙江三地游客占比 {m['yrd_rate']}%，周边市场仍有深挖空间。")
    add_bullet_cn(doc, f"文本情感比星级更严苛：SnowNLP 模型判定负面评论占 {m['model_neg_rate']}%，远高于按星级的差评率（{m['negative_rate']:.1f}%）；其中 {m['silent_neg']} 条高分评价被模型判为负面（隐性不满），提示需在好评中挖掘改进线索。")

    add_heading_cn(doc, "7.2 改进建议", level=2)
    add_bullet_cn(doc, "优化排队体验：在旺季增加观光车、电瓶车运力，设置预约排队系统，增设遮阳/避雨设施。")
    add_bullet_cn(doc, "提升性价比感知：推出家庭套票、二销组合（门票+喂食胡萝卜+观光车），让游客感知更超值。")
    add_bullet_cn(doc, '加强服务培训：差评中提及"工作人员"较多，应提升一线员工服务意识与沟通效率。')
    add_bullet_cn(doc, "平衡淡旺季客流：针对秋季淡季设计主题活动（如观鸟季、摄影节），通过价格杠杆分流旺季压力。")
    add_bullet_cn(doc, "拓展远程市场：在保持长三角优势的同时，通过内容营销吸引山东、安徽、北京等潜力客群。")
    add_bullet_cn(doc, "重视口碑管理：鼓励优质评价传播，对差评及时回复整改，持续监测 3 星及以下评价关键词。")
    doc.add_page_break()


def build_chapter_appendix(doc, m):
    """第八章：附录（数据采集说明 + 分析方法）。"""
    add_heading_cn(doc, "八、附录", level=1)

    add_heading_cn(doc, "8.1 数据采集说明", level=2)
    add_paragraph_cn(doc, f"数据来源于携程旅行网移动端 API（https://m.ctrip.com/restapi/soa2/13444/json/getCommentCollapseList），按最新排序采集，共采集有效评价 {m['total']} 条。由于携程 API 分页限制，部分历史评价可能未完全覆盖。")

    add_heading_cn(doc, "8.2 分析方法", level=2)
    add_bullet_cn(doc, "数据清洗：pandas 处理缺失值、解析时间字段、提取衍生特征（季节、评论长度等）。")
    add_bullet_cn(doc, "可视化：matplotlib 生成静态图表，嵌入 Word 报告。")
    add_bullet_cn(doc, "文本分析：jieba 中文分词 + 词频统计 + 词云生成。")
    add_bullet_cn(doc, "评分等级：好评 4-5 星，中评 3 星，差评 1-2 星。")
    add_bullet_cn(doc, "情感分析：引入 SnowNLP 中文情感极性模型，对每条评论做 0-1 语义打分（正面≥0.6 / 负面≤0.4 / 中性之间），补充星级评分的语义视角，详见第六章 6.5 节。")

    add_heading_cn(doc, "8.3 其他说明", level=2)
    add_paragraph_cn(doc, "报告日期：{}".format(datetime.now().strftime("%Y年%m月%d日")))


def main():
    """报告生成主入口：加载数据 -> 计算指标 -> 逐章编排 -> 保存。"""
    df, df_valid = data_utils.load_reviews()
    metrics = compute_metrics(df, df_valid)

    doc = Document()
    build_cover(doc, metrics)
    build_toc(doc)
    build_chapter_summary(doc, metrics)
    build_chapter_overview(doc, df, metrics)
    build_chapter_scores(doc, metrics)
    build_chapter_time(doc, metrics)
    build_chapter_users(doc, metrics)
    build_chapter_text(doc, metrics)
    build_chapter_findings(doc, metrics)
    build_chapter_appendix(doc, metrics)

    print(f"正在保存报告: {config.DOCX_REPORT}")
    doc.save(config.DOCX_REPORT)
    print("报告生成完成!")


if __name__ == "__main__":
    main()
